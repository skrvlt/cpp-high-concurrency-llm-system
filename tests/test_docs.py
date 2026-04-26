import unittest
from pathlib import Path
import zipfile
from docx import Document
from docx.oxml.ns import qn


class DocsTests(unittest.TestCase):
    def _unique_row_cells(self, row):
        cells = []
        seen = set()
        for cell in row.cells:
            marker = id(cell._tc)
            if marker not in seen:
                seen.add(marker)
                cells.append(cell)
        return cells

    def _border_value(self, cell, side):
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            return None
        border = borders.find(qn(f"w:{side}"))
        if border is None:
            return None
        return border.get(qn("w:val"))

    def assert_open_sided_table_borders(self, table):
        table_xml = table._tbl.xml
        self.assertNotIn('w:tblStyle w:val="TableGrid"', table_xml)
        for row in table.rows:
            cells = self._unique_row_cells(row)
            if not cells:
                continue
            self.assertEqual(self._border_value(cells[0], "left"), "nil")
            self.assertEqual(self._border_value(cells[0], "start"), "nil")
            self.assertEqual(self._border_value(cells[-1], "right"), "nil")
            self.assertEqual(self._border_value(cells[-1], "end"), "nil")

            for left_cell, right_cell in zip(cells, cells[1:]):
                self.assertEqual(self._border_value(left_cell, "right"), "single")
                self.assertEqual(self._border_value(left_cell, "end"), "single")
                self.assertEqual(self._border_value(right_cell, "left"), "single")
                self.assertEqual(self._border_value(right_cell, "start"), "single")

    def test_thesis_contains_required_sections(self):
        text = (Path.cwd() / "output" / "doc" / "毕业设计说明书初稿.md").read_text(
            encoding="utf-8"
        )
        for section in [
            "摘要",
            "第1章 绪论",
            "第3章 系统需求分析",
            "第4章 系统总体设计",
            "第5章 系统详细实现",
            "第6章 系统测试与结果分析",
            "结论",
        ]:
            self.assertIn(section, text)

    def test_thesis_mentions_core_figures_and_tables(self):
        text = (Path.cwd() / "output" / "doc" / "毕业设计说明书初稿.md").read_text(
            encoding="utf-8"
        )
        for marker in [
            "图4-1 系统总体结构图",
            "图4-2 智能问答处理流程图",
            "表4-1 用户表结构",
            "表4-2 会话表结构",
            "表6-4 并发性能测试结果表",
        ]:
            self.assertIn(marker, text)

    def test_thesis_mentions_sequence_and_er_guidance(self):
        text = (Path.cwd() / "output" / "doc" / "毕业设计说明书初稿.md").read_text(
            encoding="utf-8"
        )
        for marker in [
            "问答处理时序图",
            "系统 E-R 图",
            "附录B 后续补充材料说明",
        ]:
            self.assertIn(marker, text)

    def test_docs_mention_windows_linux_wsl_support(self):
        runbook = (Path.cwd() / "docs" / "runbook.md").read_text(encoding="utf-8")
        platform_doc = (Path.cwd() / "docs" / "platform-support.md").read_text(
            encoding="utf-8"
        )
        thesis = (Path.cwd() / "output" / "doc" / "毕业设计说明书初稿.md").read_text(
            encoding="utf-8"
        )
        for text in [runbook, platform_doc, thesis]:
            self.assertIn("Windows", text)
            self.assertIn("Linux", text)
            self.assertIn("WSL", text)

    def test_docs_mention_health_contract_and_env_examples(self):
        readme = (Path.cwd() / "README.md").read_text(encoding="utf-8")
        runbook = (Path.cwd() / "docs" / "runbook.md").read_text(encoding="utf-8")
        platform_doc = (Path.cwd() / "docs" / "platform-support.md").read_text(
            encoding="utf-8"
        )
        thesis = (Path.cwd() / "output" / "doc" / "毕业设计说明书初稿.md").read_text(
            encoding="utf-8"
        )
        for text in [readme, runbook, platform_doc, thesis]:
            self.assertIn("/api/health", text)
        for text in [readme, runbook, platform_doc]:
            self.assertIn(".env.example", text)

    def test_docs_mention_wsl_build_fallback(self):
        readme = (Path.cwd() / "cpp_gateway" / "README.md").read_text(encoding="utf-8")
        runbook = (Path.cwd() / "docs" / "runbook.md").read_text(encoding="utf-8")
        validation = (Path.cwd() / "docs" / "gateway-validation.md").read_text(
            encoding="utf-8"
        )
        for text in [readme, runbook, validation]:
            self.assertIn("cmake", text)
            self.assertIn("g++", text)
            self.assertIn("validate_gateway_wsl.sh", text)
            self.assertIn("API_MODE", text)

    def test_use_case_tables_are_separated_by_body_text(self):
        lines = (Path.cwd() / "output" / "doc" / "毕业设计说明书初稿.md").read_text(
            encoding="utf-8"
        ).splitlines()
        table_indices = [
            index
            for index, line in enumerate(lines)
            if line.strip() in {
                "表3-2 用户登录用例描述",
                "表3-3 智能问答用例描述",
                "表3-4 管理员配置维护用例描述",
            }
        ]
        self.assertEqual(len(table_indices), 3)

        for left, right in zip(table_indices, table_indices[1:]):
            between = [
                line.strip()
                for line in lines[left + 1 : right]
                if line.strip()
                and line.strip() not in {"[用例表]", "[/用例表]"}
                and not line.strip().startswith("用例")
                and not line.strip().startswith("主参与者")
                and not line.strip().startswith("前置条件")
                and not line.strip().startswith("后置条件")
                and not line.strip().startswith("基本事件流")
                and not line.strip().startswith("异常事件流")
                and not line.strip()[0:2].isdigit()
            ]
            self.assertTrue(
                between,
                "相邻用例描述表之间缺少说明正文",
            )

    def test_generated_doc_use_case_tables_remove_outer_borders(self):
        docx_path = Path.cwd() / "output" / "doc" / "毕业设计说明书初稿.docx"
        doc = Document(docx_path)
        for table_index in [1, 3, 5]:
            self.assert_open_sided_table_borders(doc.tables[table_index])

    def test_generated_doc_use_case_tables_keep_inner_vertical_lines(self):
        docx_path = Path.cwd() / "output" / "doc" / "毕业设计说明书初稿.docx"
        doc = Document(docx_path)
        for table_index in [1, 2, 3, 4, 5, 6]:
            self.assert_open_sided_table_borders(doc.tables[table_index])

    def test_generated_doc_table_3_1_removes_outer_borders(self):
        docx_path = Path.cwd() / "output" / "doc" / "毕业设计说明书初稿.docx"
        doc = Document(docx_path)
        self.assert_open_sided_table_borders(doc.tables[0])

    def test_midterm_docx_tables_3_1_to_3_4_remove_outer_borders(self):
        docx_paths = [
            Path.cwd() / "中期检查" / "毕业设计前三章-中期检查版.docx",
            Path.cwd() / "中期检查" / "毕业设计前三章-中期检查版-格式整理副本.docx",
        ]
        for docx_path in docx_paths:
            doc = Document(docx_path)
            for table_index in range(7):
                with self.subTest(docx=docx_path.name, table_index=table_index):
                    self.assert_open_sided_table_borders(doc.tables[table_index])
