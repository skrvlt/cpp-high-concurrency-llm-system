from pathlib import Path
import re
import shutil
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from doc_build_utils import (
    clear_document_body,
    configure_thesis_page,
    render_markdownish_document,
    set_run_font,
)
from thesis_figure_builder import ensure_all_figures_generated


ROOT = Path(__file__).resolve().parents[1]
MIDTERM_DIR = ROOT / "中期检查"
TEMPLATE = MIDTERM_DIR / "2022级毕业设计（论文）中期检查报告模板.docx"
FORMAT_TEMPLATE = MIDTERM_DIR / "毕业设计说明书（论文）格式模板参考（试行）——加页眉.docx"
THESIS_MD = ROOT / "output" / "doc" / "毕业设计说明书初稿.md"
MIDTERM_REPORT = MIDTERM_DIR / "2022级毕业设计（论文）中期检查报告-已填写.docx"
MIDTERM_THESIS_MD = MIDTERM_DIR / "毕业设计前三章-中期检查版.md"
MIDTERM_THESIS_DOCX = MIDTERM_DIR / "毕业设计前三章-中期检查版.docx"
MIDTERM_FIGURES_DIR = MIDTERM_DIR / "figures"
HEADER_IMAGE = MIDTERM_DIR / "school_header.png"


def replace_cell_text(cell, text):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for index, line in enumerate(text.splitlines()):
        run = paragraph.add_run(line)
        set_run_font(run, size=11, bold=False)
        if index != len(text.splitlines()) - 1:
            run.add_break()


def generate_midterm_report():
    doc = Document(str(TEMPLATE))
    table = doc.tables[0]

    table.cell(0, 1).text = "周殿波"
    table.cell(0, 3).text = "软件22-1班"
    table.cell(0, 5).text = "2022024763"

    table.cell(1, 1).text = "高辉"
    table.cell(1, 3).text = "讲师"
    table.cell(1, 5).text = "校内 √  校外 □"

    table.cell(2, 1).text = "基于C++高并发架构与语言大模型的交互系统的开发"

    student_summary = (
        "一、学生设计（论文）工作总结\n"
        "对照开题报告中的时间进程，截至目前已基本完成中期阶段的主要任务，整体进度与开题计划基本一致。\n"
        "1. 已完成毕业设计说明书前三章的撰写与修改，包括第1章绪论、第2章相关技术与理论基础、第3章系统需求分析。其中，绪论部分补充了研究背景、国内外研究现状、研究内容与论文结构；相关技术部分完善了 C++ 高性能网络编程、Reactor、Epoll、线程池、FastAPI、MySQL、跨语言通信等理论内容；需求分析部分补充了功能需求、非功能需求、可行性分析、系统角色、业务流程、用例描述以及数据与接口需求分析。\n"
        "2. 已初步完成系统总体架构设计，明确采用“前端页面 + C++ 高并发接入层 + Python FastAPI 智能服务层 + MySQL 数据层”的分层结构，并形成了统一的接口和端口约定。\n"
        "3. 已完成项目代码主骨架搭建，前端页面、Python 服务、数据库脚本、C++ 网关骨架、运行脚本、跨环境文档等内容已经建立，系统主链路具备进一步联调和测试的基础。\n"
        "4. 已完成部分核心功能实现与验证，包括登录、智能问答、历史记录、管理员概览、日志与配置接口等，并补充了健康检查、运行时配置和跨环境验证脚本。\n"
        "5. 当前存在的主要问题为：WSL 与 Windows 主机之间的网络访问行为存在差异，导致 C++ 网关跨环境联调过程中需要进一步区分本机回环、主机地址和 WSL 本地部署路径。针对该问题，已通过补充 WSL 本地 API 联调脚本、环境变量配置和运行文档的方式进行修正，并为后续完整压测与截图采集做好准备。\n"
        "6. 下一阶段计划继续完成第4章至第6章内容，重点推进系统总体设计细化、详细实现说明、系统测试与结果分析，同时继续完善 WSL/Linux 下的真实联调验证、图表绘制、测试数据整理和答辩材料准备。\n"
        "总体来看，当前已按时间进程完成毕业设计中期阶段应完成的论文写作与系统开发任务，具备继续完成后续设计、说明书完善和按时参加答辩的条件。\n\n"
        "学生本人（签字）：\n"
        "2026年5月6日"
    )
    teacher_summary = (
        "二、指导教师自查情况及结论说明\n"
        "该生能够按照开题报告中的计划推进毕业设计工作，目前已完成课题前三章内容的系统撰写和较大幅度的修改完善，论文结构基本符合中期检查要求。系统实现方面，已完成总体架构方案确定、主要功能模块拆分、项目代码骨架搭建以及部分核心接口联调，阶段性成果较为明确。\n"
        "在设计推进过程中，学生能够主动分析并解决跨环境联调、系统配置、文档组织等问题，表现出较强的独立思考与工程实现能力。当前仍需在后续阶段继续加强系统详细实现说明、图表绘制、测试数据整理以及最终论文规范化表达，但总体进度正常，可控性较好。\n"
        "综合判断，该生已按进度完成中期阶段相应任务，具备继续完成毕业设计后续内容并按时参加答辩的条件。\n\n"
        "指导教师（签字）：\n"
        "2026年5月7日"
    )
    group_summary = (
        "三、检查小组意见及结论\n"
        "该课题选题符合软件工程与计算机专业培养目标，具有较好的工程实践价值。学生中期阶段已完成前三章撰写，并在需求分析、技术路线和系统架构方面形成了较完整的思路；同时已具备较明确的项目实现基础。建议后续继续加强图表规范性、测试结果分析和论文格式细节，进一步完善系统联调与实验数据支撑。\n"
        "评价结论：优 □  良 √  一般 □  较差 □  很差 □\n\n"
        "检查小组组长（签字）：\n"
        "2026年5月8日"
    )

    replace_cell_text(table.cell(3, 0), student_summary)
    replace_cell_text(table.cell(4, 0), teacher_summary)
    replace_cell_text(table.cell(5, 0), group_summary)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=11, bold=False)

    doc.save(str(MIDTERM_REPORT))


def _extract_section(text: str, start: str, end: str | None = None) -> str:
    if start not in text:
        raise ValueError(f"无法找到章节起始标记：{start}")
    section = text.split(start, 1)[1]
    section = start + section
    if end and end in section:
        section = section.split(end, 1)[0]
    return section.strip()


def _renumber_chapter(section: str, from_chapter: int, to_chapter: int, new_title: str) -> str:
    section = re.sub(
        rf"## 第{from_chapter}章 .+",
        f"## 第{to_chapter}章 {new_title}",
        section,
        count=1,
    )
    section = re.sub(rf"### {from_chapter}\.", f"### {to_chapter}.", section)
    section = re.sub(rf"#### {from_chapter}\.", f"#### {to_chapter}.", section)
    section = section.replace(f"表{from_chapter}-", f"表{to_chapter}-")
    section = section.replace(f"图{from_chapter}-", f"图{to_chapter}-")
    section = section.replace(f"UC0{from_chapter}-", f"UC0{to_chapter}-")
    section = section.replace(f"第{from_chapter}章", f"第{to_chapter}章")
    return section


def _build_chapter_one() -> str:
    return """## 第1章 绪论

### 1.1 研究背景与意义

近年来，以 ChatGPT、DeepSeek、通义千问、智谱 GLM 等为代表的大语言模型快速发展，推动自然语言处理技术由传统规则驱动和浅层机器学习阶段迈向以预训练模型为核心的智能化阶段。与传统问答系统相比，大语言模型在上下文理解、复杂语义表达、开放域知识组织以及多轮对话生成方面表现更加突出，已逐渐应用于智能客服、办公助手、知识库问答、教育辅导和行业咨询等场景。

但是，从软件系统工程角度看，模型能力的提升并不等同于应用系统成熟度的提升。一个面向真实使用场景的智能交互平台不仅要能够生成自然语言回答，还需要具备稳定的网络接入能力、清晰的请求处理链路、可维护的数据记录机制以及必要的后台管理能力。若直接采用单体式脚本服务承载浏览器接入、请求解析、模型调用、历史记录和管理功能，在多用户同时访问时容易暴露连接管理粗糙、资源回收不及时、请求链路阻塞和日志追踪困难等问题。

高并发接入问题在 Web 应用、即时通信、在线考试和智能服务平台中具有共性。C++ 在底层网络开发、系统调用控制和运行效率方面具有优势，适合承担连接监听、I/O 多路复用和请求转发等基础工作；Python 在大模型接口封装、业务逻辑组织和快速开发方面生态成熟，适合承担智能服务层任务。因此，将 C++ 高并发接入层与 Python 大模型服务层进行分层协同，是兼顾性能、开发效率和可维护性的一种可行方案。

本课题以“基于 C++ 高并发架构与语言大模型的交互系统的开发”为研究对象，设计并实现一套面向浏览器问答场景的智能交互系统。系统通过前端页面完成用户交互，通过 C++ 网关完成高并发连接接入与请求转发，通过 Python FastAPI 服务完成身份校验、上下文组织、模型调用和结果返回，并通过数据库保存用户、会话、消息、配置和日志等数据。该方案既体现高性能网络编程能力，又能结合大语言模型应用开发，具有较强的工程实践意义。

本研究的意义主要体现在以下三个方面。第一，在理论层面，探索高性能网络架构与大语言模型服务化部署的结合方式，为“底层高性能接入 + 上层智能能力封装”的系统设计提供分析样本。第二，在工程层面，验证 C++ 接入层与 Python AI 服务协同设计的可行性，为类似智能交互系统的模块划分、接口组织和运行验证提供参考。第三，在实践层面，形成一套适合本科毕业设计实施的完整方案，包括需求分析、概要设计、系统实现、测试验证和论文撰写支撑。

### 1.2 国内外研究现状

在国外研究与工程实践中，高性能服务器领域较早开展了事件驱动网络模型、I/O 多路复用机制和高并发服务端编程研究。Epoll、Kqueue、Reactor、Proactor 等模型在互联网服务端开发中已有广泛应用，相关实践表明，事件驱动模型在高连接数场景下通常比简单的一连接一线程模型具有更好的资源利用率和响应能力。与此同时，FastAPI、gRPC、异步消息队列、容器化部署和模型网关等技术也被广泛用于智能服务系统建设，使大模型能力能够以接口形式接入实际业务系统。

在人工智能服务部署方面，国外研究较早关注模型服务化、推理加速和 AI 基础设施优化。随着大模型应用从研究验证走向业务落地，系统架构不再只关注模型效果本身，还需要考虑请求调度、接口稳定性、上下游服务隔离、异常降级和部署复现等工程问题。对于智能问答类应用而言，如何将模型调用能力放入一个稳定、可扩展、可监控的服务体系中，已经成为大模型应用工程化的重要方向。

国内方面，随着大模型生态快速发展，越来越多的系统开始尝试将语言模型能力接入企业办公、知识问答、智能客服和教育辅助场景。学术界和产业界一方面关注模型本身的训练方法和能力提升，另一方面也更加重视模型服务的工程化组织。围绕问答系统、知识库问答、智能体协同和业务辅助决策等方向，已经出现较多基于语言模型的应用研究与项目实践。

在底层服务实现上，C++ 仍然是高性能网络系统的重要选择；在模型调用和业务封装层面，Python 由于生态完善、开发效率高，成为大模型服务开发的常用语言。因此，采用“C++ 负责高并发接入，Python 负责智能推理”的分层方案具有较好的现实基础。对于本科毕业设计而言，该方案能够同时覆盖前后端协同、数据库设计、网络编程、接口封装和系统测试等多个核心能力点，避免课题内容过于单一。

现有研究和实践仍存在一定不足：一类工作更关注底层网络性能与框架设计，但对大模型调用和业务封装涉及较少；另一类工作更关注模型效果和应用体验，但较少讨论高并发接入场景下的连接管理、资源回收和故障处理。二者结合后的工程化实现细节仍有进一步梳理空间。本课题正是在这一背景下，尝试构建一套规模可控、功能闭环明确、技术路线清晰的智能交互系统。

### 1.3 主要研究内容与论文组织结构

围绕课题目标，本文主要完成以下研究内容：第一，分析高并发智能交互系统的业务特点、用户角色、功能需求和非功能需求；第二，设计前端页面、C++ 高并发接入层、Python FastAPI 智能服务层和数据库层构成的系统总体架构；第三，设计用户登录、智能问答、历史记录、后台概览、日志查看和配置维护等核心功能模块；第四，设计用户、会话、消息、日志和配置等数据库表结构；第五，设计系统接口、业务流程和跨环境运行方案，为后续实现与测试提供依据。

在技术路线上，本文以浏览器端问答场景为牵引，从需求分析出发，逐步形成“Web 前端 - C++ 接入层 - Python 智能服务层 - MySQL 数据层”的四层结构。浏览器通过 HTTP 发起登录、问答、历史记录和后台管理请求；C++ 网关负责连接管理、协议解析和请求转发；Python 服务负责用户校验、上下文组织、模型调用、日志记录和数据返回；数据库负责对关键业务数据进行持久化存储。系统同时支持 Windows 下的前端与 Python 服务直连演示，以及 Linux / WSL 下经 C++ 网关转发的双后端协同演示。

本课题拟重点解决以下问题：高并发接入层与智能业务层的职责边界如何划分；多轮问答、历史记录、配置管理和日志追踪等功能如何形成统一业务闭环；C++ 网关与 Python 服务之间如何通过 HTTP + JSON 建立清晰接口契约；系统如何在 Windows、Linux、WSL 等不同环境下给出明确运行路径，降低只能在单一设备演示的风险。

全文组织结构如下。第1章为绪论，介绍研究背景与意义、国内外研究现状、主要研究内容和论文组织结构。第2章为需求分析，分析系统功能需求、非功能需求、可行性、关键业务流程和用例。第3章为概要设计，给出系统总体架构、功能模块、数据库结构、接口设计和界面设计。后续章节将在完整论文中继续展开系统详细实现、测试分析、结论、参考文献和附录等内容。"""


def _insert_design_additions(chapter: str) -> str:
    addition = """
### 4.6 接口设计

系统接口采用 HTTP + JSON 方式组织，前端、C++ 网关与 Python 服务之间通过统一路径和统一数据格式完成通信。接口设计遵循职责清晰、参数明确、错误可识别的原则，使前端页面能够在直连 Python 服务和通过 C++ 网关转发两种模式下复用同一套调用逻辑。

核心接口设计如表4-5所示。
表4-5 核心接口设计表

| 接口路径 | 请求方法 | 主要参数 | 返回内容 | 说明 |
| --- | --- | --- | --- | --- |
| /api/health | GET | 无 | 服务状态、运行模式、模型名称 | 用于运行检查与部署验证 |
| /api/login | POST | username、password | token、角色、会话编号 | 用户登录与身份建立 |
| /api/chat | POST | token、message | 回复内容、会话标题 | 智能问答主接口 |
| /api/history | GET | token | 历史问答列表 | 查询当前用户会话历史 |
| /api/admin/overview | GET | token | 用户数、会话数、消息数、日志数 | 管理员系统概览 |
| /api/admin/logs | GET | token | 日志列表 | 管理员查看运行日志 |
| /api/admin/config | GET/POST | token、config_key、config_value | 当前配置或更新结果 | 管理员配置维护 |

C++ 网关在接口设计中承担接入层角色，对浏览器请求进行 HTTP 解析并转发至 Python 服务。Python 服务承担业务接口实现角色，负责 token 校验、权限判断、上下文组织和返回数据封装。对于身份无效、权限不足、参数错误和上游服务不可达等情况，系统通过 HTTP 状态码和 JSON 错误信息返回结果，避免前端只能依赖文本判断异常。

### 4.7 界面设计

系统界面采用浏览器访问方式，主要包括普通用户端和管理员端两类页面。普通用户端由登录区域、问答区域和历史记录区域组成：登录区域用于输入账号密码并展示当前运行模式；问答区域用于输入问题、显示多轮对话内容和当前会话标题；历史记录区域用于查看当前用户已保存的问答记录。该设计能够满足中期阶段“登录 - 提问 - 返回答案 - 保存历史”的核心演示链路。

管理员端页面主要包含系统概览、日志查看和配置维护三部分。系统概览展示用户数、会话数、消息数、日志数和模型名称等指标；日志查看用于观察登录、问答和配置修改等事件；配置维护用于调整模型名称、超时参数等运行配置。管理员页面通过服务端权限校验控制访问，避免只依赖前端页面隐藏入口。

前端界面在运行模式上支持直连 Python 服务和经 C++ 网关转发两种方式。用户可通过 URL 参数或运行时配置切换访问路径，从而在 Windows 环境下快速演示通用功能，在 WSL / Linux 环境下展示 C++ 高并发网关参与请求转发的系统结构。"""
    marker = "### 4.6 跨环境部署设计"
    if marker not in chapter:
        raise ValueError("无法在概要设计章节中定位跨环境部署设计小节")
    return chapter.replace(marker, addition + "\n\n### 4.8 跨环境部署设计", 1)


def build_midterm_thesis_lines():
    text = THESIS_MD.read_text(encoding="utf-8")
    abstract = text.split("## 目录", 1)[0].strip()
    abstract = re.sub(r"^# .+\n+", "", abstract)
    chapter_two = _renumber_chapter(
        _extract_section(text, "## 第3章 系统需求分析", "## 第4章 系统总体设计"),
        3,
        2,
        "系统需求分析",
    )
    chapter_three_source = _extract_section(text, "## 第4章 系统总体设计", "## 第5章 系统详细实现")
    chapter_three_source = _insert_design_additions(chapter_three_source)
    chapter_three = _renumber_chapter(chapter_three_source, 4, 3, "系统概要设计")

    toc = """## 目录

第1章 绪论 1
1.1 研究背景与意义 1
1.2 国内外研究现状 2
1.3 主要研究内容与论文组织结构 3
第2章 系统需求分析 4
2.1 系统目标 4
2.2 可行性分析 4
2.3 功能需求分析 5
2.4 非功能需求分析 7
2.5 系统角色分析 8
2.6 业务流程分析 11
2.7 数据与接口需求分析 12
2.8 本章小结 13
第3章 系统概要设计 14
3.1 系统设计原则 14
3.2 系统总体架构 14
3.3 功能模块设计 15
3.4 数据库设计 16
3.5 系统流程设计 18
3.6 接口设计 20
3.7 界面设计 21
3.8 跨环境部署设计 22

Contents
Chapter 1 Introduction 1
1.1 Research Background and Significance 1
1.2 Research Status at Home and Abroad 2
1.3 Main Research Contents and Thesis Organization 3
Chapter 2 System Requirements Analysis 4
2.1 System Objectives 4
2.2 Feasibility Analysis 4
2.3 Functional Requirements Analysis 5
2.4 Non-functional Requirements Analysis 7
2.5 System Role Analysis 8
2.6 Business Process Analysis 11
2.7 Data and Interface Requirements Analysis 12
2.8 Summary 13
Chapter 3 System Outline Design 14
3.1 Design Principles 14
3.2 Overall System Architecture 14
3.3 Functional Module Design 15
3.4 Database Design 16
3.5 System Process Design 18
3.6 Interface Design 20
3.7 Interface Page Design 21
3.8 Cross-environment Deployment Design 22"""

    document = "\n\n".join(
        [
            "[分页]",
            abstract,
            "[分页]",
            toc,
            "[分页]",
            _build_chapter_one(),
            "[分页]",
            chapter_two,
            "[分页]",
            chapter_three,
        ]
    )
    clean_lines = [
        "# 基于C++高并发架构与语言大模型的交互系统的开发",
        "",
    ] + [line for line in document.splitlines() if line.strip() != "[分页]"]
    MIDTERM_THESIS_MD.write_text("\n".join(clean_lines).rstrip() + "\n", encoding="utf-8")
    return document.splitlines()


def prepare_midterm_figures():
    ensure_all_figures_generated(MIDTERM_FIGURES_DIR)
    aliases = {
        "图3-1": "图2-1",
        "图3-2": "图2-2",
        "图3-3": "图2-3",
        "图4-1": "图3-1",
        "图4-2": "图3-2",
        "图4-3": "图3-3",
        "图4-4": "图3-4",
    }
    for source, target in aliases.items():
        source_path = MIDTERM_FIGURES_DIR / f"{source}.png"
        target_path = MIDTERM_FIGURES_DIR / f"{target}.png"
        shutil.copyfile(source_path, target_path)


def _ensure_header_image():
    if HEADER_IMAGE.exists():
        return
    with ZipFile(FORMAT_TEMPLATE) as template_zip:
        HEADER_IMAGE.write_bytes(template_zip.read("word/media/image1.png"))


def _add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def _apply_header_footer(doc: Document):
    _ensure_header_image()
    for section in doc.sections:
        header = section.header
        footer = section.footer
        header.is_linked_to_previous = False
        footer.is_linked_to_previous = False

        header_para = header.paragraphs[0]
        header_para.text = ""
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.paragraph_format.space_before = 0
        header_para.paragraph_format.space_after = 0
        header_run = header_para.add_run()
        header_run.add_picture(str(HEADER_IMAGE), width=Cm(15.8))

        footer_para = footer.paragraphs[0]
        footer_para.text = ""
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_field(footer_para)
        for run in footer_para.runs:
            set_run_font(run, size=10.5, bold=False, font_name="宋体")


def generate_midterm_thesis_docx(lines):
    doc = Document(str(FORMAT_TEMPLATE)) if FORMAT_TEMPLATE.exists() else Document()
    clear_document_body(doc)
    configure_thesis_page(doc)
    _apply_header_footer(doc)
    prepare_midterm_figures()
    render_markdownish_document(
        doc,
        lines,
        "基于C++高并发架构与语言大模型的交互系统的开发",
        figures_dir=MIDTERM_FIGURES_DIR,
    )
    doc.save(str(MIDTERM_THESIS_DOCX))


def main():
    MIDTERM_DIR.mkdir(parents=True, exist_ok=True)
    generate_midterm_report()
    lines = build_midterm_thesis_lines()
    generate_midterm_thesis_docx(lines)


if __name__ == "__main__":
    main()
