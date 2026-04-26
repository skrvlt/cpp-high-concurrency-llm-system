from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


@dataclass
class UseCaseBlock:
    caption: str
    case_id: str = ""
    name: str = ""
    description: str = ""
    actor: str = ""
    precondition: str = ""
    postcondition: str = ""
    basic_flow: list[str] = field(default_factory=list)
    exception_flow: list[str] = field(default_factory=list)


FIGURE_CAPTION_RE = re.compile(r"^(图\d+-\d+)\s+(.+)$")


def clean_inline_markdown(text: str) -> str:
    return (
        text.replace("**", "")
        .replace("__", "")
        .replace("`", "")
        .replace("[用例表]", "")
        .replace("[/用例表]", "")
        .strip()
    )


def set_run_font(run, size=12, bold=False, font_name="宋体", east_asia_font=None):
    east_asia = east_asia_font or font_name
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def configure_thesis_page(doc: Document):
    """Apply the school thesis template page size and margins."""
    for section in doc.sections:
        section.page_width = Mm(184)
        section.page_height = Mm(260)
        section.top_margin = Mm(45)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(23)
        section.right_margin = Mm(17)
        section.header_distance = Mm(15)
        section.footer_distance = Mm(15)
        sect_pr = section._sectPr
        pg_sz = sect_pr.find(qn("w:pgSz"))
        if pg_sz is None:
            pg_sz = OxmlElement("w:pgSz")
            sect_pr.append(pg_sz)
        pg_sz.set(qn("w:w"), "10433")
        pg_sz.set(qn("w:h"), "14742")

        pg_mar = sect_pr.find(qn("w:pgMar"))
        if pg_mar is None:
            pg_mar = OxmlElement("w:pgMar")
            sect_pr.append(pg_mar)
        for attr, value in {
            "top": "2552",
            "right": "964",
            "bottom": "851",
            "left": "1304",
            "header": "851",
            "footer": "851",
            "gutter": "0",
        }.items():
            pg_mar.set(qn(f"w:{attr}"), value)

        doc_grid = sect_pr.find(qn("w:docGrid"))
        if doc_grid is None:
            doc_grid = OxmlElement("w:docGrid")
            sect_pr.append(doc_grid)
        doc_grid.set(qn("w:type"), "linesAndChars")
        doc_grid.set(qn("w:linePitch"), "343")
        doc_grid.set(qn("w:charSpace"), "1434")


def clear_document_body(doc: Document):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_paragraph(doc: Document, text: str, style: str = "body"):
    display_text = {"摘要": "摘 要", "目录": "目 录"}.get(text, text)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24) if style == "body" else Pt(0)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(0)
    if style == "title":
        p.style = doc.styles["Title"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(display_text)
        set_run_font(run, size=18, bold=True, font_name="黑体")
    elif style == "h1":
        p.style = doc.styles["Heading 1"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(17)
        p.paragraph_format.space_after = Pt(16.5)
        p.paragraph_format.line_spacing = 2.408333333333333
        run = p.add_run(display_text)
        if display_text == "Abstract":
            set_run_font(run, size=18, bold=True, font_name="Times New Roman", east_asia_font="黑体")
        else:
            set_run_font(run, size=18, bold=False, font_name="黑体")
    elif style == "h2":
        p.style = doc.styles["Heading 2"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(13)
        p.paragraph_format.space_after = Pt(13)
        p.paragraph_format.line_spacing = 1.7333333333333334
        run = p.add_run(display_text)
        set_run_font(run, size=15, bold=False, font_name="黑体")
    elif style == "h3":
        p.style = doc.styles["Heading 3"]
        p.paragraph_format.space_before = Pt(13)
        p.paragraph_format.space_after = Pt(13)
        p.paragraph_format.line_spacing = 1.7333333333333334
        run = p.add_run(display_text)
        set_run_font(run, size=14, bold=False, font_name="黑体")
    else:
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(clean_inline_markdown(display_text))
        set_run_font(run, size=12, bold=False)
    return p


def add_figure(doc: Document, image_path: Path, caption_text: str):
    picture_paragraph = doc.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.first_line_indent = Pt(0)
    picture_paragraph.paragraph_format.space_before = Pt(0)
    picture_paragraph.paragraph_format.space_after = Pt(0)
    run = picture_paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(15.5))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.line_spacing = Pt(20)
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(0)
    run = caption.add_run(caption_text)
    set_run_font(run, size=10.5, bold=False, font_name="黑体")


def _set_cell_text(cell, text: str, bold: bool = False, align="center"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_border(cell, side: str, value: str | None = None, size: int = 4):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    border = tc_borders.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        tc_borders.append(border)

    if value is None:
        border.set(qn("w:val"), "nil")
        border.attrib.pop(qn("w:sz"), None)
        border.attrib.pop(qn("w:space"), None)
        border.attrib.pop(qn("w:color"), None)
        return

    border.set(qn("w:val"), value)
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "auto")


def _set_table_border(table, side: str, value: str | None = None, size: int = 4):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    border = tbl_borders.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        tbl_borders.append(border)

    if value is None:
        border.set(qn("w:val"), "nil")
        border.attrib.pop(qn("w:sz"), None)
        border.attrib.pop(qn("w:space"), None)
        border.attrib.pop(qn("w:color"), None)
        return

    border.set(qn("w:val"), value)
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "auto")


def _remove_outer_vertical_borders(table):
    for side in ("left", "right", "start", "end"):
        _set_table_border(table, side, None)
    for row in table.rows:
        if not row.cells:
            continue
        _set_cell_border(row.cells[0], "left", None)
        _set_cell_border(row.cells[-1], "right", None)


def _iter_unique_cells(row):
    seen = set()
    cells = []
    for cell in row.cells:
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        cells.append(cell)
    return cells


def _apply_use_case_table_borders(table):
    _remove_outer_vertical_borders(table)
    for row in table.rows:
        unique_cells = _iter_unique_cells(row)
        for cell in unique_cells:
            for side in ("top", "bottom"):
                _set_cell_border(cell, side, "single")
            for side in ("left", "right"):
                _set_cell_border(cell, side, None)

        if len(unique_cells) >= 2:
            _set_cell_border(unique_cells[0], "right", "single")
            _set_cell_border(unique_cells[1], "left", "single")
        if len(unique_cells) >= 3:
            _set_cell_border(unique_cells[1], "right", "single")
            _set_cell_border(unique_cells[2], "left", "single")


def add_use_case_table(doc: Document, block: UseCaseBlock):
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(block.caption)
    set_run_font(run, size=10.5, bold=False, font_name="黑体")
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=0, cols=3)

    simple_rows = [
        ("用例编号", block.case_id),
        ("用例名称", block.name),
        ("用例概述", block.description),
        ("主参与者", block.actor),
        ("前置条件", block.precondition),
        ("后置条件", block.postcondition),
    ]
    for label, value in simple_rows:
        row = table.add_row()
        _set_cell_text(row.cells[0], label)
        merged = row.cells[1].merge(row.cells[2])
        _set_cell_text(merged, value, align="left")

    flow_sections = [
        ("基本事件流", block.basic_flow),
        ("异常事件流", block.exception_flow),
    ]
    for section_name, items in flow_sections:
        if not items:
            continue
        header_row = table.add_row()
        _set_cell_text(header_row.cells[1], "步骤")
        _set_cell_text(header_row.cells[2], "活动")
        data_rows = []
        for index, item in enumerate(items, start=1):
            row = table.add_row()
            _set_cell_text(row.cells[1], str(index))
            _set_cell_text(row.cells[2], item, align="left")
            data_rows.append(row)
        merged = header_row.cells[0]
        for row in data_rows:
            merged = merged.merge(row.cells[0])
        _set_cell_text(merged, section_name)

    _apply_use_case_table_borders(table)

    if block.exception_flow:
        ext_table = doc.add_table(rows=0, cols=3)
        header = ext_table.add_row()
        _set_cell_text(header.cells[0], "扩展事件流")
        _set_cell_text(header.cells[1], "步骤")
        _set_cell_text(header.cells[2], "活动")
        for index, item in enumerate(block.exception_flow, start=1):
            row = ext_table.add_row()
            _set_cell_text(row.cells[0], "")
            _set_cell_text(row.cells[1], f"{index}a")
            _set_cell_text(row.cells[2], item, align="left")
        _apply_use_case_table_borders(ext_table)


def parse_use_case_block(caption: str, raw_lines: list[str]) -> UseCaseBlock:
    block = UseCaseBlock(caption=caption)
    current_section: str | None = None
    key_map = {
        "用例编号": "case_id",
        "用例名称": "name",
        "用例描述": "description",
        "用例概述": "description",
        "主参与者": "actor",
        "前置条件": "precondition",
        "后置条件": "postcondition",
    }

    for raw in raw_lines:
        line = raw.strip()
        if not line:
            continue
        if line.endswith("：") or line.endswith(":"):
            section = line[:-1]
            if section in {"基本事件流", "异常事件流"}:
                current_section = section
                continue
        if current_section and line[0].isdigit() and "." in line:
            content = line.split(".", 1)[1].strip()
            if current_section == "基本事件流":
                block.basic_flow.append(content)
            else:
                block.exception_flow.append(content)
            continue

        current_section = None
        sep = "：" if "：" in line else ":"
        if sep in line:
            key, value = [part.strip() for part in line.split(sep, 1)]
            attr = key_map.get(key)
            if attr:
                setattr(block, attr, value)
    return block


def _is_markdown_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|")
    if not stripped:
        return False
    return all(part.strip().replace("-", "").replace(":", "") == "" for part in stripped.split("|"))


def _parse_markdown_row(line: str) -> list[str]:
    return [clean_inline_markdown(part.strip()) for part in line.strip().strip("|").split("|")]


def add_markdown_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row_data in enumerate(rows):
        row = table.add_row()
        for cell, value in zip(row.cells, row_data):
            _set_cell_text(cell, value, bold=(row_index == 0), align="center" if row_index == 0 else "left")


def _add_toc_entry(doc: Document, text: str, page: str, level: int = 1, english: bool = False):
    p = doc.add_paragraph()
    try:
        p.style = doc.styles[f"toc {min(max(level, 1), 3)}"]
    except KeyError:
        pass
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt((level - 1) * 18)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(14.2),
        WD_TAB_ALIGNMENT.RIGHT,
        WD_TAB_LEADER.DOTS,
    )
    run = p.add_run(f"{text}\t{page}")
    if english:
        set_run_font(run, size=12, bold=(level == 1), font_name="Times New Roman", east_asia_font="宋体")
    else:
        set_run_font(run, size=12, bold=False, font_name="宋体")


def _toc_line_parts(line: str) -> tuple[str, str, int] | None:
    if "\t" in line:
        title, page = line.rsplit("\t", 1)
    else:
        match = re.match(r"^(.+?)\s+(\d+|[IVXLCDM]+)$", line)
        if not match:
            return None
        title, page = match.group(1), match.group(2)
    title = title.strip()
    page = page.strip()
    if title.startswith("第") or title.startswith("Chapter") or title in {"结 论", "致 谢", "参考文献"}:
        level = 1
    else:
        dot_count = title.split(" ", 1)[0].count(".")
        level = 3 if dot_count >= 2 else 2
    return title, page, level


def render_markdownish_document(
    doc: Document,
    lines: list[str],
    title: str,
    figures_dir: Path | None = None,
):
    add_paragraph(doc, title, "title")

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line in {"[分页]", "[PAGE_BREAK]"}:
            doc.add_page_break()
            i += 1
            continue

        next_index = i + 1
        while next_index < len(lines) and not lines[next_index].strip():
            next_index += 1
        next_line = lines[next_index].strip() if next_index < len(lines) else ""
        if line.startswith("表") and "用例描述" in line and next_line == "[用例表]":
            i = next_index + 1
            block_lines = []
            while i < len(lines) and lines[i].strip() != "[/用例表]":
                block_lines.append(lines[i])
                i += 1
            add_use_case_table(doc, parse_use_case_block(line, block_lines))
            i += 1
            continue

        if line.startswith("|") and next_index < len(lines) and _is_markdown_table_separator(lines[next_index]):
            table_rows = [_parse_markdown_row(line)]
            i = next_index + 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(_parse_markdown_row(lines[i].strip()))
                i += 1
            add_markdown_table(doc, table_rows)
            continue

        if line == "## 目录":
            add_paragraph(doc, "目录", "h1")
            i += 1
            english_toc = False
            while i < len(lines):
                current = lines[i].strip()
                if current.startswith("## ") and current != "## 目录":
                    break
                if not current:
                    i += 1
                    continue
                if current == "Contents":
                    doc.add_page_break()
                    add_paragraph(doc, "Contents", "h1")
                    english_toc = True
                    i += 1
                    continue
                parts = _toc_line_parts(current)
                if parts:
                    title, page, level = parts
                    _add_toc_entry(doc, title, page, level=level, english=english_toc)
                else:
                    add_paragraph(doc, current, "body")
                i += 1
            continue

        figure_match = FIGURE_CAPTION_RE.match(line)
        if figure_match and figures_dir is not None:
            figure_id = figure_match.group(1)
            figure_path = figures_dir / f"{figure_id}.png"
            if figure_path.exists():
                add_figure(doc, figure_path, line)
                i += 1
                continue

        if line.startswith("# "):
            add_paragraph(doc, line[2:].strip(), "title")
        elif line.startswith("## "):
            add_paragraph(doc, line[3:].strip(), "h1")
        elif line.startswith("### "):
            add_paragraph(doc, line[4:].strip(), "h2")
        elif line.startswith("#### "):
            add_paragraph(doc, line[5:].strip(), "h3")
        else:
            add_paragraph(doc, line, "body")
        i += 1
